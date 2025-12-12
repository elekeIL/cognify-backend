"""Enhanced file processing service with high performance and accuracy."""

import asyncio
import logging
import re
import time
from pathlib import Path
from typing import Tuple

import aiofiles
import fitz  # PyMuPDF
import magic  # pip install python-magic (or python-magic-bin on Windows)
from docx import Document as DocxDocument

from app.models.document import FileType

# Configure logging
logger = logging.getLogger(__name__)

# Clean up excessive newlines
CLEANUP_REGEX = re.compile(r"\n{3,}")

# Allowed MIME types (prevent malicious uploads)
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

# Mapping from MIME to FileType (more reliable than extension)
MIME_TO_FILETYPE = {
    "application/pdf": FileType.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
    "text/plain": FileType.TXT,
    "application/octet-stream": FileType.TXT,  # Often used for .txt
}


def _detect_file_type_safely(file_path: Path) -> FileType:
    """Detect file type using libmagic (prevents extension spoofing)."""
    mime_type = magic.from_file(str(file_path), mime=True)
    if mime_type not in ALLOWED_MIME_TYPES:
        raise ValueError(f"Disallowed MIME type: {mime_type}")
    file_type = MIME_TO_FILETYPE.get(mime_type)
    if not file_type:
        # Fallback to extension if MIME is ambiguous
        ext = file_path.suffix.lower().lstrip(".")
        ext_mapping = {"pdf": FileType.PDF, "docx": FileType.DOCX, "txt": FileType.TXT}
        file_type = ext_mapping.get(ext)
        if not file_type:
            raise ValueError(f"Unsupported file type: {ext}")
    return file_type


class FileProcessor:
    """High-performance, accurate text extraction service."""

    @staticmethod
    async def extract_text(file_path: str | Path, original_filename: str = "") -> Tuple[str, int]:
        """
        Extract text from a file with high accuracy and performance.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Security: validate MIME type
        file_type = _detect_file_type_safely(file_path)

        logger.info(f"Extracting text from {file_path.name} ({file_type.value})")

        extractors = {
            FileType.PDF: FileProcessor._extract_from_pdf,
            FileType.DOCX: FileProcessor._extract_from_docx,
            FileType.TXT: FileProcessor._extract_from_txt,
        }

        extractor = extractors[file_type]
        text = await extractor(file_path)
        text = CLEANUP_REGEX.sub("\n\n", text).strip()

        # Validate extracted content
        if not text:
            raise ValueError(f"No text content could be extracted from {file_path.name}")

        word_count = len(text.split())

        # Log extraction results
        logger.info(f"Successfully extracted {word_count} words from {file_path.name}")

        # Warn if very low word count (may indicate extraction issues)
        if word_count < 10:
            logger.warning(f"Very low word count ({word_count}) from {file_path.name} - may indicate extraction issues")

        return text, word_count

    # ──────────────────────────────────────────────────────────────
    # PDF Extraction – FAST using PyMuPDF
    # ──────────────────────────────────────────────────────────────
    @staticmethod
    async def _extract_from_pdf(file_path: Path) -> str:
        def sync_extract():
            start_time = time.perf_counter()
            try:
                raw_texts = []
                doc = fitz.open(file_path)
                total_pages = len(doc)
                logger.info(f"Extracting text from {total_pages} pages using PyMuPDF...")

                for page_num in range(total_pages):
                    try:
                        page = doc[page_num]
                        text = page.get_text("text")

                        if text and text.strip():
                            raw_texts.append(text.strip())
                    except Exception as page_err:
                        logger.warning(f"Error on page {page_num + 1}: {page_err}")
                        continue

                doc.close()

                if not raw_texts:
                    raise ValueError("No text could be extracted from PDF")

                elapsed = time.perf_counter() - start_time
                logger.info(f"PDF extraction completed: {total_pages} pages in {elapsed:.2f}s ({total_pages/max(elapsed, 0.01):.1f} pages/sec)")

                return "\n\n".join(raw_texts)

            except Exception as e:
                logger.exception(f"PDF extraction failed: {e}")
                raise

        return await asyncio.to_thread(sync_extract)

    # ──────────────────────────────────────────────────────────────
    # DOCX Extraction – Now includes tables
    # ──────────────────────────────────────────────────────────────
    @staticmethod
    async def _extract_from_docx(file_path: Path) -> str:
        def run_in_thread():
            doc = DocxDocument(file_path)
            text_parts = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        return await asyncio.to_thread(run_in_thread)

    # ──────────────────────────────────────────────────────────────
    # TXT Extraction – Fully async
    # ──────────────────────────────────────────────────────────────
    @staticmethod
    async def _extract_from_txt(file_path: Path) -> str:
        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = await f.read()
            return content.replace("\r\n", "\n")

    # ──────────────────────────────────────────────────────────────
    # Utility: Validate extension (optional extra check)
    # ──────────────────────────────────────────────────────────────
    @staticmethod
    def validate_extension(filename: str, allowed: Tuple[str, ...] = (".pdf", ".docx", ".txt")) -> bool:
        return filename.lower().endswith(allowed)

    @staticmethod
    def validate_file_extension(filename: str, allowed_extensions: list) -> bool:
        """Check if file extension is allowed."""
        ext = filename.lower().split(".")[-1]
        return ext in allowed_extensions

    @staticmethod
    def get_file_type(filename: str) -> FileType:
        """Determine file type from filename extension."""
        ext = filename.lower().split(".")[-1]
        mapping = {
            "pdf": FileType.PDF,
            "docx": FileType.DOCX,
            "txt": FileType.TXT,
        }
        file_type = mapping.get(ext)
        if not file_type:
            raise ValueError(f"Unsupported file extension: {ext}")
        return file_type

